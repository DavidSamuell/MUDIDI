"""Mathpix Convert artifacts → flat Stage-1 transcript (adapter v1)."""

from __future__ import annotations

import logging
from pathlib import Path

from dictextractor.evaluation.stage1.normalize_typography import normalize_line
from dictextractor.ocr.adapters.layout_to_transcript_v1 import FlatTranscriptParts
from dictextractor.ocr.adapters.markdown_to_flat import find_markdown_source
from dictextractor.ocr.adapters.mathpix_lines import mathpix_transcript_from_lines_json

logger = logging.getLogger(__name__)

_MATHPIX_DOCX_NAME = "mathpix.docx"
_MATHPIX_LINES_NAME = "mathpix.lines.json"


def mathpix_docx_path(page_dir: Path) -> Path:
    """Return the canonical Mathpix DOCX artifact path under a page directory."""
    return page_dir / _MATHPIX_DOCX_NAME


def mathpix_lines_path(page_dir: Path) -> Path:
    """Return the canonical Mathpix lines.json artifact path under a page directory."""
    return page_dir / _MATHPIX_LINES_NAME


def mathpix_page_is_complete(page_dir: Path, *, stem: str) -> bool:
    """True when markdown source and flat output both exist with content."""
    flat = page_dir / f"{stem}_stage1_flat.txt"
    if not flat.is_file():
        return False
    if not flat.read_text(encoding="utf-8", errors="replace").strip():
        return False
    return find_markdown_source(page_dir, stem=stem) is not None


def _cell_lines(text: str) -> list[str]:
    """Split a table cell into normalized non-empty lines."""
    lines: list[str] = []
    for raw in text.splitlines():
        line = normalize_line(raw.strip())
        if line:
            lines.append(line)
    if not lines and text.strip():
        line = normalize_line(text.strip())
        if line:
            lines.append(line)
    return lines


def _column_major_table_lines(table) -> list[str]:
    """
    Emit flat lines in column-major order (spec v2 for multi-column tables).

    Each table column is read top-to-bottom before advancing to the next column.
    Empty cells are skipped; intra-cell newlines become separate lines.
    """
    if not table.rows:
        return []

    num_cols = max(len(row.cells) for row in table.rows)
    merged: list[str] = []
    for col_idx in range(num_cols):
        for row in table.rows:
            if col_idx >= len(row.cells):
                continue
            merged.extend(_cell_lines(row.cells[col_idx].text))
    return merged


def _paragraph_lines(doc) -> list[str]:
    """Paragraph-only Mathpix exports: preserve Mathpix paragraph order."""
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        raw = paragraph.text.strip()
        if not raw:
            continue
        for part in raw.splitlines():
            line = normalize_line(part.strip())
            if line:
                lines.append(line)
    return lines


def mathpix_transcript_from_docx(docx_path: Path) -> FlatTranscriptParts:
    """
    Build flat transcript parts from a Mathpix Convert ``.docx``.

    Table exports (e.g. trilingual Circassian) use column-major line order.
    Paragraph exports keep Mathpix paragraph order in the body (DOCX fallback).
    """
    from docx import Document

    doc = Document(str(docx_path))
    if doc.tables:
        body = _column_major_table_lines(doc.tables[0])
        for extra_table in doc.tables[1:]:
            body.extend(_column_major_table_lines(extra_table))
        if not body and doc.paragraphs:
            logger.warning(
                "Mathpix %s: tables empty; falling back to paragraphs",
                docx_path.name,
            )
            body = _paragraph_lines(doc)
    else:
        body = _paragraph_lines(doc)

    return FlatTranscriptParts(header=[], body=body, footer=[])


def mathpix_transcript_from_page_dir(page_dir: Path) -> FlatTranscriptParts:
    """Load Mathpix flat parts, preferring ``mathpix.lines.json`` when present."""
    lines_path = mathpix_lines_path(page_dir)
    if lines_path.is_file():
        try:
            return mathpix_transcript_from_lines_json(lines_path)
        except ValueError as exc:
            logger.warning(
                "Mathpix lines.json unusable in %s (%s); falling back to docx",
                page_dir,
                exc,
            )

    docx_path = mathpix_docx_path(page_dir)
    if not docx_path.is_file():
        raise FileNotFoundError(
            f"No Mathpix lines.json or docx under {page_dir}"
        )
    return mathpix_transcript_from_docx(docx_path)
