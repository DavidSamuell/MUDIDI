"""
File I/O helpers: reading source documents and writing extraction outputs.
"""

import csv
import json
from pathlib import Path
from typing import Any, Dict, List, Optional


# ---------------------------------------------------------------------------
# Readers
# ---------------------------------------------------------------------------

def read_docx_text(docx_path: str) -> str:
    """
    Read and return all non-empty text from a .docx file.

    Extracts body paragraphs and table cell text (Mathpix exports often use
    tables only, with no paragraph body).

    Args:
        docx_path: Path to the DOCX file.

    Returns:
        Newline-joined paragraph and table-row text (tab-separated cells).
    """
    from docx import Document

    doc = Document(docx_path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def read_text_file(text_path: str) -> str:
    """
    Read a plain text or Markdown file.

    Args:
        text_path: Path to the text file.

    Returns:
        File contents as a string.
    """
    with open(text_path, "r", encoding="utf-8") as f:
        return f.read()


def load_tsv(filepath: str) -> List[Dict[str, str]]:
    """
    Load a TSV file into a list of row dictionaries.

    Args:
        filepath: Path to the TSV file.

    Returns:
        List of dicts, one per row.
    """
    entries = []
    with open(filepath, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f, delimiter="\t")
        for row in reader:
            entries.append(dict(row))
    return entries


# ---------------------------------------------------------------------------
# Writers
# ---------------------------------------------------------------------------

CANONICAL_TSV_FIELDS = [
    "Block_ID",
    "Entry_Type",
    "Headword",
    "Parent_Lexeme",
    "Sense_Number",
    "Homonym_Number",
    "Gloss",
    "Gloss_Secondary",
    "Usage_Note",
    "POS",
    "Phonetic",
    "Cross_References",
    "Examples",
    "Example_Glosses",
]


def _key_to_column(key: str) -> str:
    """snake_case extra-field key → Title_Case_Underscore TSV header."""
    return "_".join(part.capitalize() for part in key.split("_") if part)


def _collect_extra_keys(entries: List[Dict]) -> List[str]:
    """
    Union of all extra_fields keys across entries, in first-appearance order.
    Skips keys whose Title_Case form would collide with a canonical column.
    """
    canonical = set(CANONICAL_TSV_FIELDS)
    seen: Dict[str, None] = {}
    for entry in entries:
        for k in (entry.get("extra_fields") or {}).keys():
            if not k or k in seen:
                continue
            if _key_to_column(k) in canonical:
                continue
            seen[k] = None
    return list(seen.keys())


def save_to_json(dictionary_page, output_path: str) -> None:
    """
    Write a DictionaryPage to a JSON file (array of entry dicts).

    Args:
        dictionary_page: DictionaryPage instance.
        output_path: Destination file path.
    """
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(
            [entry.model_dump() for entry in dictionary_page.entries],
            f,
            ensure_ascii=False,
            indent=2,
        )
    print(f"Saved JSON to {output_path}")


def _join_list_field(value) -> str:
    """Join list values for TSV; pass through strings unchanged."""
    if isinstance(value, list):
        return " | ".join(str(v) for v in value if v)
    return str(value) if value else ""


def json_to_tsv(json_path: str, output_path: Optional[str] = None) -> str:
    """
    Render a TSV file from a saved entries-JSON file.

    The TSV contains MDF-oriented canonical columns plus any discovered
    ``extra_fields`` columns (snake_case key → Title_Case header).

    Args:
        json_path: Path to the input JSON file.
        output_path: Optional destination path (defaults to same stem + .tsv).

    Returns:
        Path to the written TSV file.
    """
    json_path = Path(json_path)
    if output_path is None:
        output_path = json_path.with_suffix(".tsv")
    else:
        output_path = Path(output_path)

    with open(json_path, "r", encoding="utf-8") as f:
        entries = json.load(f)

    from dictextractor.utils.mdf_export import block_ids_for_entries

    block_ids = block_ids_for_entries(entries)
    extra_keys = _collect_extra_keys(entries)
    extra_columns = [_key_to_column(k) for k in extra_keys]
    fieldnames = CANONICAL_TSV_FIELDS + extra_columns

    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, delimiter="\t")
        writer.writeheader()
        for idx, entry in enumerate(entries):
            extras = entry.get("extra_fields") or {}
            row = {
                "Block_ID": block_ids[idx] if idx < len(block_ids) else "",
                "Entry_Type": entry.get("entry_type") or "main",
                "Headword": entry.get("headword") or entry.get("headword_phrase", ""),
                "Parent_Lexeme": entry.get("parent_lexeme", ""),
                "Sense_Number": (
                    "" if entry.get("sense_number") is None else str(entry["sense_number"])
                ),
                "Homonym_Number": (
                    "" if entry.get("homonym_number") is None else str(entry["homonym_number"])
                ),
                "Gloss": entry.get("gloss", ""),
                "Gloss_Secondary": entry.get("gloss_secondary", ""),
                "Usage_Note": entry.get("usage_note", ""),
                "POS": entry.get("pos", ""),
                "Phonetic": entry.get("phonetic", ""),
                "Cross_References": _join_list_field(entry.get("cross_references")),
                "Examples": _join_list_field(entry.get("examples")),
                "Example_Glosses": _join_list_field(entry.get("example_glosses")),
            }
            for key, column in zip(extra_keys, extra_columns):
                row[column] = extras.get(key, "")
            writer.writerow(row)

    print(f"Rendered {len(entries)} entries → {output_path} ({len(extra_columns)} extra columns)")
    return str(output_path)


def save_stage2_outputs(
    dictionary_page,
    output_json: str,
    *,
    languages_config=None,
) -> Dict[str, Any]:
    """
    Normalise, validate, and write Stage-2 JSON, TSV, MDF, and validation report.

    Returns:
        Dict with paths and validation summary.
    """
    from dictextractor.schemas.entry import DictionaryPage
    from dictextractor.utils.mdf_export import (
        entries_to_mdf_text,
        log_validation_report,
        normalize_stage2_entries,
        validate_stage2_entries,
    )

    json_path = Path(output_json)
    tsv_path = json_path.with_suffix(".tsv")
    mdf_path = json_path.with_suffix(".mdf.txt")
    validation_path = json_path.with_name(json_path.stem + "_validation.json")

    normalized = normalize_stage2_entries(dictionary_page.entries)
    page = DictionaryPage(
        entries=normalized,
        page_number=dictionary_page.page_number,
        source_file=dictionary_page.source_file,
    )
    report = validate_stage2_entries(normalized, languages_config)
    log_validation_report(report, page_label=json_path.stem)

    save_to_json(page, str(json_path))
    json_to_tsv(str(json_path), str(tsv_path))
    mdf_path.write_text(
        entries_to_mdf_text(normalized, languages_config),
        encoding="utf-8",
    )
    validation_path.write_text(
        json.dumps(
            {
                "ok": report.ok,
                "errors": [
                    {"message": i.message, "entry_index": i.entry_index, "block_id": i.block_id}
                    for i in report.errors
                ],
                "warnings": [
                    {"message": i.message, "entry_index": i.entry_index, "block_id": i.block_id}
                    for i in report.warnings
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    print(f"Saved MDF to {mdf_path}")
    print(f"Saved validation report to {validation_path}")
    return {
        "json": str(json_path),
        "tsv": str(tsv_path),
        "mdf": str(mdf_path),
        "validation": str(validation_path),
        "validation_ok": report.ok,
        "error_count": len(report.errors),
        "warning_count": len(report.warnings),
    }
